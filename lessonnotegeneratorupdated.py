import os
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from docx import Document
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from groq import Groq
from together import Together
import re

load_dotenv()

class LessonNoteGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Avalon STEM Lesson Note Generator")
        self.root.geometry("900x700")
        self.root.configure(bg="#ffffff") # Keep root background white or primary_bg

        # Initialize the container
        self.container = ttk.Frame(self.root, padding=20)
        self.container.pack(fill='both', expand=True)

        # STEM subjects list
        self.stem_subjects = [
            "mathematics", "maths", "further mathematics", "further maths",
            "chemistry", "physics", "biology", 
            "agricultural science", "computer science", "geography"
        ]

        # Configure style
        self.style = ttk.Style(self.root)
        self.style.theme_use('clam') # 'clam' is a good base for customization

        # Set fonts and colors for modern minimalist look
        self.heading_font = ("Segoe UI Semibold", 24, 'bold') # Adjusted for the new heading
        self.label_font = ("Segoe UI", 12) # Slightly smaller for minimalism
        self.entry_font = ("Segoe UI", 12)
        self.text_font = ("Segoe UI", 12)
        self.button_font = ("Segoe UI Semibold", 13) # Slightly smaller
        
        # Modern Minimalist Palette
        self.primary_bg = "#f8f9fa" # Light grey background for main areas
        self.secondary_bg = "#e9ecef" # Slightly darker for frames/cards
        self.primary_text_color = "#343a40" # Dark grey for main text
        self.secondary_text_color = "#6c757d" # Muted grey for secondary text (not heavily used here, but good to define)
        self.accent_color = "#007bff" # Blue for primary actions
        self.border_color = "#dee2e6" # Light border for subtle separation

        # Apply styles
        self.style.configure('TFrame', background=self.primary_bg)
        self.style.configure('Card.TFrame', background=self.secondary_bg, relief='flat', borderwidth=1, bordercolor=self.border_color) # For scrollable_frame
        
        # Configure TLabel to inherit background from parent or explicitly set
        self.style.configure('TLabel', background=self.secondary_bg, foreground=self.primary_text_color, font=self.label_font)
        
        self.style.configure('TEntry', fieldbackground="#ffffff", foreground=self.primary_text_color, font=self.entry_font, borderwidth=1, relief='solid', bordercolor=self.border_color)
        self.style.configure('TCombobox', fieldbackground="#ffffff", foreground=self.primary_text_color, font=self.entry_font, borderwidth=1, relief='solid', bordercolor=self.border_color)
        
        # TLabelframe styling
        self.style.configure('TLabelframe', background=self.secondary_bg, foreground=self.primary_text_color, font=self.label_font, borderwidth=1, relief='solid', bordercolor=self.border_color)
        self.style.configure('TLabelframe.Label', background=self.secondary_bg, foreground=self.primary_text_color, font=self.label_font) # For the label of the labelframe

        # Primary Button Styling
        self.style.configure('Primary.TButton',
                             background=self.accent_color,
                             foreground="#ffffff", # White text on accent color
                             font=self.button_font,
                             borderwidth=0, # Flat button
                             relief='flat',
                             padding=(15, 8)) # More padding for a modern feel
        self.style.map('Primary.TButton',
                       background=[('active', '#0056b3'), ('pressed', '#0056b3')], # Darker blue on hover/active
                       foreground=[('active', '#ffffff'), ('pressed', '#ffffff')])

        # Setup UI
        self.setup_ui()

    def setup_ui(self):
        outer_frame = ttk.Frame(self.container, padding=20) # Increased padding
        outer_frame.pack(fill='both', expand=True)

        # Add the main heading at the top
        heading_label = ttk.Label(outer_frame, text="Avalon Lesson Note Generator", font=self.heading_font,
                                  background=self.primary_bg, foreground=self.primary_text_color)
        heading_label.pack(pady=(0, 20)) # Add some padding below the heading

        # Canvas & scrollbar for scrollable content area
        self.canvas = tk.Canvas(outer_frame, borderwidth=0, highlightthickness=0, background=self.primary_bg) # Use primary_bg
        scrollbar = ttk.Scrollbar(outer_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas, style='Card.TFrame') # Apply Card.TFrame style

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Layout parameters
        pad_x_label = 12 # Increased padding
        pad_y_label = 8
        pad_x_entry = 12 # Increased padding
        pad_y_entry = 8

        # Week
        ttk.Label(self.scrollable_frame, text="Week:", style='TLabel').grid(row=0, column=0, sticky='w', padx=pad_x_label, pady=pad_y_label)
        self.week_entry = ttk.Entry(self.scrollable_frame, width=30, style='TEntry')
        self.week_entry.grid(row=0, column=1, sticky='ew', padx=pad_x_entry, pady=pad_y_entry)

        # Class
        ttk.Label(self.scrollable_frame, text="Class:", style='TLabel').grid(row=1, column=0, sticky='w', padx=pad_x_label, pady=pad_y_label)
        self.class_var = tk.StringVar()
        class_values = [
            "Nursery 1", "Nursery 2", "Nursery 3",
            "Primary 1", "Primary 2", "Primary 3",
            "Primary 4", "Primary 5", "Primary 6",
            "JSS 1", "JSS 2", "JSS 3",
            "SSS 1", "SSS 2", "SSS 3"
        ]
        self.class_cb = ttk.Combobox(self.scrollable_frame, textvariable=self.class_var, state="readonly",
                                    values=class_values, width=30, style='TCombobox')
        self.class_cb.grid(row=1, column=1, sticky='ew', padx=pad_x_entry, pady=pad_y_entry)

        # Subject
        ttk.Label(self.scrollable_frame, text="Subject:", style='TLabel').grid(row=2, column=0, sticky='w', padx=pad_x_label, pady=pad_y_label)
        self.subject_entry = ttk.Entry(self.scrollable_frame, width=30, style='TEntry')
        self.subject_entry.grid(row=2, column=1, sticky='ew', padx=pad_x_entry, pady=pad_y_entry)

        # Topic
        ttk.Label(self.scrollable_frame, text="Topic:", style='TLabel').grid(row=3, column=0, sticky='w', padx=pad_x_label, pady=pad_y_label)
        self.topic_entry = ttk.Entry(self.scrollable_frame, width=30, style='TEntry')
        self.topic_entry.grid(row=3, column=1, sticky='ew', padx=pad_x_entry, pady=pad_y_entry)

        # Behavioral Objectives
        obj_frame = ttk.LabelFrame(self.scrollable_frame, text="Behavioral Objectives (Minimum 3)", padding=14, style='TLabelframe')
        obj_frame.grid(row=4, column=0, columnspan=2, sticky='ew', padx=20, pady=(20, 36))
        obj_frame.columnconfigure(1, weight=1)

        self.objective_entries = []
        for i in range(5):
            # Labels within LabelFrame should also use the TLabel style
            label = ttk.Label(obj_frame, text=f"{i + 1}.", style='TLabel')
            label.grid(row=i, column=0, sticky='nw', padx=6, pady=6)
            entry = ttk.Entry(obj_frame, style='TEntry')
            entry.grid(row=i, column=1, sticky='ew', padx=6, pady=6)
            self.objective_entries.append(entry)

        # Generate Button
        self.generate_btn = ttk.Button(self.scrollable_frame, text="Generate Lesson Note",
                                      command=self.generate_note, style='Primary.TButton')
        self.generate_btn.grid(row=5, column=0, columnspan=2, pady=(0, 26), sticky='ew', padx=20) # Added padx

        # Output Preview
        output_label = ttk.Label(self.scrollable_frame, text="Lesson Note Preview:",
                                style='TLabel') # Apply TLabel style
        output_label.grid(row=6, column=0, columnspan=2, sticky='w', padx=pad_x_label, pady=(0, 12))

        self.output_text = scrolledtext.ScrolledText(self.scrollable_frame,
                                                   height=18,
                                                   font=self.text_font,
                                                   wrap='word',
                                                   relief='solid', # Subtle solid border
                                                   bd=1, # 1 pixel border
                                                   background='#ffffff', # White background for text area
                                                   foreground=self.primary_text_color,
                                                   highlightbackground=self.border_color, # Border color when not focused
                                                   highlightcolor=self.accent_color, # Accent color when focused
                                                   highlightthickness=1) # 1 pixel highlight
        self.output_text.grid(row=7, column=0, columnspan=2, sticky='nsew', padx=20, pady=(0, 20))

        # Export Button
        self.export_btn = ttk.Button(self.scrollable_frame, text="Export Lesson Note (DOCX)",
                                    command=self._export_docx, style='Primary.TButton')
        self.export_btn.grid(row=8, column=0, columnspan=2, pady=(0, 20), sticky='ew', padx=20) # Added padx

        # Grid configuration
        self.scrollable_frame.columnconfigure(1, weight=1)
        self.scrollable_frame.rowconfigure(7, weight=1)

    def is_stem_subject(self, subject):
        return subject.lower() in self.stem_subjects

    def generate_note(self):
        try:
            inputs = {
                'week': self.week_entry.get(),
                'class': self.class_var.get(),
                'subject': self.subject_entry.get(),
                'topic': self.topic_entry.get(),
                'objectives': [entry.get() for entry in self.objective_entries if entry.get()],
            }

            if len(inputs['objectives']) < 3:
                messagebox.showerror("Error", "Please enter at least 3 objectives")
                return

            # Generate content with STEM-specific handling
            generated_steps = []
            for obj in inputs['objectives']:
                step_content = self.clean_ai_response(self.call_ai_api(obj, inputs['subject']))
                generated_steps.append(step_content)

            evaluation_questions = self.generate_evaluation_questions(inputs['topic'], inputs['objectives'], inputs['subject'])
            assignment_questions = self.generate_assignment_questions(inputs['topic'], inputs['objectives'], inputs['subject'])
            image_notice = self.check_for_image_requirements(inputs['topic'], inputs['objectives'], inputs['subject'])
            
            # Generate key formulae/equations for STEM subjects
            key_formulae = ""
            if self.is_stem_subject(inputs['subject']):
                key_formulae = self.generate_key_formulae(inputs['topic'], inputs['subject'])

            complete_inputs = {
                **inputs,
                'generated_steps': generated_steps,
                'evaluation_questions': evaluation_questions,
                'assignment_questions': assignment_questions,
                'image_notice': image_notice,
                'is_stem': self.is_stem_subject(inputs['subject']),
                'key_formulae': key_formulae
            }

            self.lesson_note = self.build_template(complete_inputs)
            self.output_text.delete(1.0, tk.END)
            self.output_text.insert(tk.END, self.lesson_note)

        except Exception as e:
            messagebox.showerror("Error", f"Generation failed: {str(e)}")

    def clean_ai_response(self, text):
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Remove common introductory/concluding phrases and markdown asterisks
            cleaned_line = line.replace('**', '').strip() # Remove asterisks and strip whitespace
            if not any(phrase.lower() in cleaned_line.lower() for phrase in [
                "here are", "based on", "let me know",
                "step", "here is", "meets your requirements",
                "finally", "in summary", "as requested",
                "i hope this", "please note", "additional notes",
                "in conclusion", "to summarize", "in brief", "overall"
            ]):
                if cleaned_line: # Only add non-empty lines
                    cleaned_lines.append(cleaned_line)
        return '\n'.join(cleaned_lines).strip()


    def call_ai_api(self, objective, subject):
        if self.is_stem_subject(subject):
            prompt = self.get_stem_prompt(objective, subject)
        else:
            prompt = f"""
            Generate concise content directly addressing this objective: {objective}
            - Include any recommended visual aids within the step content as: [Insert image showing...]
            - Be direct and factual
            - Remove all introductory phrases and concluding remarks.
            - Ensure proper spacing and line breaks for readability.
            - If defining, provide exactly 2 definitions.
            - If explaining, provide clear steps or points.
            - Format as plain text with no markdown (e.g., no asterisks for bolding).
            """
            
        try:
            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
            response = client.chat.completions.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=4000
            )
            content = response.choices[0].message.content
            
            # Post-process STEM content
            if self.is_stem_subject(subject):
                content = self.format_stem_content(content, subject)
                
            return self.clean_ai_response(content)
        except Exception as e:
            print(f"Error calling Groq API: {e}")
            return self.call_together_ai_api(objective, subject)

    def get_stem_prompt(self, objective, subject):
        """Generate specialized prompt for STEM subjects"""
        subject_lower = subject.lower()
        
        # Common instructions for all STEM prompts
        common_instructions = """
            - Be direct and factual with no introductory phrases or concluding remarks.
            - Ensure proper spacing and line breaks for readability.
            - Format as plain text with no markdown (e.g., no asterisks for bolding).
            """

        if subject_lower in ["mathematics", "maths", "further mathematics", "further maths"]:
            return f"""
            Generate concise mathematical content for this objective: {objective}
            - Use proper Unicode mathematical notation (e.g., π, √, ∫, ∑, θ, ≠, ≤, ≥).
            - For equations, use proper formatting (e.g., x² + y² = z²).
            - Include key formulas where relevant.
            - Provide step-by-step solutions for problems.
            - Include recommended visual aids within the step content as: [Insert diagram showing...].
            {common_instructions}
            """
        elif subject_lower == "physics":
            return f"""
            Generate concise physics content for this objective: {objective}
            - Use proper physics notation (e.g., Δx, F=ma, μ, λ).
            - Include relevant formulas with units.
            - Use proper Unicode symbols (e.g., Ω, °, ±, →).
            - Include recommended visual aids within the step content as: [Insert diagram showing...].
            {common_instructions}
            """
        elif subject_lower == "chemistry":
            return f"""
            Generate concise chemistry content for this objective: {objective}
            - Use proper chemical notation (e.g., H₂O, CO₂, CH₄).
            - For equations, use proper arrow symbols (→, ⇌).
            - Include state symbols where appropriate (s, l, g, aq).
            - Use proper Unicode symbols (e.g., °C, ΔH, λ).
            - Include recommended visual aids within the step content as: [Insert diagram showing...].
            {common_instructions}
            """
        elif subject_lower == "biology":
            return f"""
            Generate concise biology content for this objective: {objective}
            - Use proper biological terminology.
            - Include key processes with clear steps.
            - Use proper notation for species names (e.g., Homo sapiens).
            - Include recommended visual aids within the step content as: [Insert diagram showing...].
            {common_instructions}
            """
        elif subject_lower == "geography":
            return f"""
            Generate concise geography content for this objective: {objective}
            - Use proper geographical terminology.
            - Include key concepts and processes with clear explanations.
            - Include recommended visual aids within the step content as: [Insert map showing...].
            {common_instructions}
            """
        else:
            return f"""
            Generate concise content directly addressing this objective: {objective}
            - Include recommended visual aids within the step content as: [Insert image showing...].
            - Be direct and factual.
            - Remove all introductory phrases and concluding remarks.
            - Ensure proper spacing and line breaks for readability.
            - If defining, provide exactly 2 definitions.
            - If explaining, provide clear steps or points.
            - Format as plain text with no markdown (e.g., no asterisks for bolding).
            """

    def format_stem_content(self, text, subject):
        """Post-process STEM content to ensure proper formatting"""
        subject_lower = subject.lower()
        
        # Common replacements for all STEM subjects
        replacements = {
            'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ',
            'theta': 'θ', 'pi': 'π', 'sigma': 'σ', 'omega': 'ω',
            '->': '→', '=>': '⇒', 'sqrt': '√', 'integral': '∫',
            'sum': '∑', 'product': '∏', 'infinity': '∞',
            '!=': '≠', '<=': '≤', '>=': '≥', '+-': '±',
            'deg': '°', 'lambda': 'λ', 'ohm': 'Ω',
            'approx': '≈', 'plusminus': '±', 'times': '×', 'divide': '÷'
        }
        
        for plain, unicode in replacements.items():
            text = text.replace(plain, unicode)
        
        # Subject-specific replacements
        if subject_lower in ["chemistry", "physics", "biology"]:
            text = text.replace('_2', '₂').replace('_3', '₃').replace('_4', '₄').replace('_5', '₅')
            text = text.replace('^+', '⁺').replace('^-', '⁻').replace('^2+', '²⁺').replace('^2-', '²⁻')
            
        if subject_lower == "chemistry":
            text = text.replace('<->', '⇌').replace('<=>', '⇌')
            
        # Format mathematical equations
        if subject_lower in ["mathematics", "physics"]:
            # Handle exponents (x^2 → x²)
            text = re.sub(r'(\w)\^(\d+)', lambda m: m.group(1) + ''.join('⁰¹²³⁴⁵⁶⁷⁸⁹'[int(d)] for d in m.group(2)), text)
            # Handle subscripts (x_1 → x₁)
            text = re.sub(r'(\w)_(\d+)', lambda m: m.group(1) + ''.join('₀₁₂₃₄₅₆₇₈₉'[int(d)] for d in m.group(2)), text)
            # Handle simple fractions like a/b to a⁄b (using fraction slash)
            text = re.sub(r'(\d+)/(\d+)', r'\1⁄\2', text)
            
        return text

    def call_together_ai_api(self, objective, subject):
        prompt = self.get_stem_prompt(objective, subject) if self.is_stem_subject(subject) else f"""
        Generate concise content directly addressing this objective: {objective}
        - Include recommended visual aids within the step content as: [Insert image showing...].
        - Be direct and factual.
        - Remove all introductory phrases and concluding remarks.
        - Ensure proper spacing and line breaks for readability.
        - If defining, provide exactly 2 definitions.
        - If explaining, provide clear steps or points.
        - Format as plain text with no markdown (e.g., no asterisks for bolding).
        """
        try:
            client = Together(api_key=os.getenv("TOGETHER_AI_API_KEY"))
            response = client.chat.completions.create(
                model="together-model",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=4000
            )
            content = response.choices[0].message.content
            
            # Post-process STEM content
            if self.is_stem_subject(subject):
                content = self.format_stem_content(content, subject)
                
            return self.clean_ai_response(content)
        except Exception as e:
            print(f"Error calling Together API: {e}")
            return "[Error generating lesson step]"

    def generate_evaluation_questions(self, topic, objectives, subject):
        num_questions = len(objectives)
        prompt = f"""
        Generate {num_questions} evaluation questions for topic '{topic}' based on these objectives: {', '.join(objectives)}.
        - Questions should directly test each objective.
        - Format as a numbered list.
        - No introductory or concluding phrases.
        - Ensure proper spacing and line breaks.
        """
        return self.clean_ai_response(self.call_groq_api(prompt))

    def generate_assignment_questions(self, topic, objectives, subject):
        prompt = f"""
        Generate 2-3 relevant assignment questions or tasks for topic '{topic}' based on these objectives: {', '.join(objectives)}.
        - The questions/tasks should encourage deeper understanding and application of the lesson.
        - Format as a numbered list.
        - No introductory or concluding phrases.
        - Ensure proper spacing and line breaks.
        """
        return self.clean_ai_response(self.call_groq_api(prompt))

    def generate_key_formulae(self, topic, subject):
        """Generates key formulae/equations for STEM subjects."""
        subject_lower = subject.lower()
        if subject_lower in ["mathematics", "physics", "chemistry"]:
            prompt = f"""
            Generate 3-5 key formulae or equations relevant to the {subject} topic '{topic}'.
            - Use proper Unicode mathematical/chemical notation.
            - Present as a numbered list.
            - Do not include any introductory or concluding phrases.
            - Ensure proper spacing and line breaks.
            """
            return self.clean_ai_response(self.call_groq_api(prompt))
        return ""

    def check_for_image_requirements(self, topic, objectives, subject):
        if not self.is_stem_subject(subject):
            return ""
            
        prompt = f"""
        Analyze if teaching this {subject} topic '{topic}' with these objectives {', '.join(objectives)} would require visual aids/images.
        If images are needed, list specific image types that would be helpful for teaching this lesson.
        - Provide a maximum of 4 distinct image types.
        - List them as a numbered or bulleted list.
        - If no images are needed, respond with "No specific visual aids recommended for this topic."
        - Do not include any introductory or concluding phrases.
        """
        response = self.call_groq_api(prompt).strip()
        # Ensure the response is clean and doesn't contain unwanted phrases
        cleaned_response = self.clean_ai_response(response)
        if cleaned_response and "no specific visual aids" not in cleaned_response.lower():
            return f"Recommended visual aids:\n{cleaned_response}"
        return "" # Return empty string if no specific aids or if the AI explicitly says none

    def call_groq_api(self, prompt):
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=4000
        )
        return response.choices[0].message.content

    def build_template(self, inputs):
        if inputs['is_stem']:
            lesson_note = f"STEM LESSON NOTE\n\n"
            lesson_note += f"Class: {inputs['class']}\n"
            lesson_note += f"Week: {inputs['week']}\n"
            lesson_note += f"Subject: {inputs['subject']}\n"
            lesson_note += f"Topic: {inputs['topic']}\n\n"
            
            if inputs['key_formulae']:
                lesson_note += "KEY FORMULAE/EQUATIONS:\n"
                lesson_note += f"{inputs['key_formulae']}\n\n"
            else:
                lesson_note += "KEY FORMULAE/EQUATIONS:\n"
                lesson_note += "[No specific key formulae for this subject/topic]\n\n"
        else:
            lesson_note = f"LESSON NOTE\n\n"
            lesson_note += f"Class: {inputs['class']}\n"
            lesson_note += f"Week: {inputs['week']}\n"
            lesson_note += f"Subject: {inputs['subject']}\n"
            lesson_note += f"Topic: {inputs['topic']}\n\n"

        lesson_note += "BEHAVIORAL OBJECTIVES:\n"
        for i, obj in enumerate(inputs['objectives'], 1):
            lesson_note += f"{i}. {obj}\n"

        lesson_note += "\nPRESENTATION STEPS:\n"
        for i, step in enumerate(inputs['generated_steps'], 1):
            lesson_note += f"Step {i}: {step}\n"

        lesson_note += "\nSTUDENTS ACTIVITIES:\n"
        lesson_note += "Students listen attentively, participate in discussions, ask questions, and take notes.\n"

        lesson_note += "\nEVALUATION:\n"
        lesson_note += f"{inputs['evaluation_questions']}\n"

        lesson_note += "\nSUMMARY:\n"
        lesson_note += "The teacher summarizes the key points of the lesson.\n"

        lesson_note += "\nCONCLUSION:\n"
        lesson_note += "The teacher concludes the lesson and reinforces the main concepts.\n"

        lesson_note += "\nASSIGNMENT/CLASS ACTIVITY:\n"
        lesson_note += f"{inputs['assignment_questions']}\n"

        if inputs['image_notice']:
            lesson_note += f"\nIMAGE NOTICE:\n{inputs['image_notice']}\n"

        return lesson_note

    def _get_base_filename(self, extension):
        """Helper to generate a clean base filename."""
        class_name = self.class_var.get()
        subject = self.subject_entry.get().strip()
        topic = self.topic_entry.get().strip()
        
        def clean_text_for_filename(text):
            if not text:
                return ""
            invalid_chars = '<>:"/\\|?*'
            for char in invalid_chars:
                text = text.replace(char, '_')
            return text.strip('. ')
        
        class_clean = clean_text_for_filename(class_name)
        subject_clean = clean_text_for_filename(subject)
        topic_clean = clean_text_for_filename(topic)
        
        filename_parts = []
        if class_clean:
            filename_parts.append(class_clean)
        if subject_clean:
            filename_parts.append(subject_clean)
        if topic_clean:
            filename_parts.append(topic_clean)
        
        if filename_parts:
            return "_".join(filename_parts) + "_Lesson_Note." + extension
        else:
            return "Lesson_Note." + extension

    def _create_docx_document_object(self):
        """Helper function to create and populate a docx Document object."""
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial Unicode MS'
        font.size = Pt(11)

        topic = self.topic_entry.get().strip()
        class_name = self.class_var.get()
        subject = self.subject_entry.get().strip()

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(f"Lesson Note on {topic}" if topic else "Lesson Note")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.alignment = 1

        # Info table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.2)

        # Standard fields
        fields = [
            ("Week", self.week_entry.get()),
            ("Date", ""),
            ("Class", class_name),
            ("Subject", subject),
            ("Topic", topic),
            ("Duration", ""),
            ("Sex", ""),
            ("Age", ""),
            ("Entry Behavior", ""),
            ("Teaching Aid", ""),
            ("Reference Text", ""),
            ("Introduction", "")
        ]

        for field, value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = value

        lesson_text = self.output_text.get("1.0", tk.END)
        
        # Extract Key Formulae
        if self.is_stem_subject(subject):
            key_formulae_text = ""
            key_formulae_start = lesson_text.find("KEY FORMULAE/EQUATIONS:")
            if key_formulae_start != -1:
                next_section_start = lesson_text.find("BEHAVIORAL OBJECTIVES:", key_formulae_start)
                if next_section_start != -1:
                    key_formulae_text = lesson_text[key_formulae_start + len("KEY FORMULAE/EQUATIONS:"):next_section_start].strip()
                else:
                    key_formulae_text = lesson_text[key_formulae_start + len("KEY FORMULAE/EQUATIONS:"):].strip()
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Key Formulae/Equations"
            row_cells[1].text = key_formulae_text if key_formulae_text else "[No specific key formulae for this subject/topic]"

        # Behavioral Objectives
        row_cells = table.add_row().cells
        row_cells[0].text = "Behavioral Objectives"
        objs = [entry.get() for entry in self.objective_entries if entry.get()]
        row_cells[1].text = "\n".join(f"{i + 1}. {obj}" for i, obj in enumerate(objs))

        # Helper to extract sections
        def extract_section(text, start_label, end_label):
            start = text.find(start_label)
            end = text.find(end_label)
            if start == -1 or end == -1 or end <= start:
                return ""
            return text[start + len(start_label):end].strip()

        # Presentation Steps
        pres_steps_text = extract_section(lesson_text, "PRESENTATION STEPS:", "STUDENTS ACTIVITIES:")
        steps = []
        current_step_lines = []
        for line in pres_steps_text.splitlines():
            if line.strip().startswith("Step"):
                if current_step_lines:
                    steps.append("\n".join(current_step_lines))
                    current_step_lines = []
            if line.strip():
                current_step_lines.append(line.strip())
        if current_step_lines:
            steps.append("\n".join(current_step_lines))

        for step in steps:
            if step:
                row_cells = table.add_row().cells
                parts = step.split(":", maxsplit=1)
                row_cells[0].text = parts[0].strip()
                row_cells[1].text = parts[1].strip() if len(parts) > 1 else ""

        # Students Activities
        students_act_text = extract_section(lesson_text, "STUDENTS ACTIVITIES:", "EVALUATION:")
        if students_act_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Students Activities"
            row_cells[1].text = students_act_text

        # Evaluation
        eval_text = extract_section(lesson_text, "EVALUATION:", "SUMMARY:")
        if eval_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Evaluation"
            row_cells[1].text = eval_text

        # Summary
        summary_text = extract_section(lesson_text, "SUMMARY:", "CONCLUSION:")
        if summary_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Summary"
            row_cells[1].text = summary_text

        # Conclusion
        conclusion_text = extract_section(lesson_text, "CONCLUSION:", "ASSIGNMENT/CLASS ACTIVITY:")
        if conclusion_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Conclusion"
            row_cells[1].text = conclusion_text

        # Assignment
        assign_text = extract_section(lesson_text, "ASSIGNMENT/CLASS ACTIVITY:", "IMAGE NOTICE:")
        if not assign_text:
            # If no image notice, get everything after assignment
            assign_start = lesson_text.find("ASSIGNMENT/CLASS ACTIVITY:")
            if assign_start != -1:
                assign_text = lesson_text[assign_start + len("ASSIGNMENT/CLASS ACTIVITY:"):].strip()
        
        if assign_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Assignment/Class Activity"
            row_cells[1].text = assign_text

        # Image Notice
        if "IMAGE NOTICE:" in lesson_text:
            row_cells = table.add_row().cells
            row_cells[0].text = "Recommended Visual Aids"
            image_notice_text = lesson_text.split("IMAGE NOTICE:")[1].strip()
            row_cells[1].text = image_notice_text

        return doc

    def _export_docx(self):
        """Exports the lesson note as a DOCX file."""
        default_filename = self._get_base_filename("docx")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            initialfile=default_filename
        )
        
        if not file_path:
            return

        try:
            doc = self._create_docx_document_object()
            doc.save(file_path)
            messagebox.showinfo("Success", "Lesson note exported successfully as DOCX!")
        except Exception as e:
            messagebox.showerror("Error", f"DOCX Export failed: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = LessonNoteGenerator(root)
    root.mainloop()
