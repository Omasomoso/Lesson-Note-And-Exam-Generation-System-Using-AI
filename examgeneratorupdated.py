import os
import re
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
from tkinter import font as tkfont
from typing import List, Optional, Tuple
from PIL import Image, ImageTk
import requests
from docx import Document
from docx.shared import Pt # Import Pt for font sizing
from dotenv import load_dotenv
from groq import Groq

# Load environment variables
load_dotenv()


class ModernButton(ttk.Button):
    """A modern styled button with hover effects"""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.style = ttk.Style()
        self.style.configure(
            "Modern.TButton",
            font=("Segoe UI", 10, "bold"),
            padding=10,
            relief="flat",
            foreground="#ffffff",
            background="#4a6da7",
            bordercolor="#4a6da7"
        )
        self.style.map(
            "Modern.TButton",
            background=[
                ('pressed', '#3a5a8f'),
                ('active', '#5b7db8'),
                ('disabled', '#cccccc')
            ],
            foreground=[
                ('pressed', '#ffffff'),
                ('active', '#ffffff'),
                ('disabled', '#888888')
            ],
            bordercolor=[
                ('pressed', '#3a5a8f'),
                ('active', '#5b7db8')
            ]
        )
        self.configure(style="Modern.TButton")


class ExamQuestionGenerator:
    """A professional GUI application for generating exam questions using AI APIs."""
    
    # Modern color palette
    PRIMARY_COLOR = "#4a6da7"  # Deep blue
    SECONDARY_COLOR = "#6dd5fa"  # Light blue
    ACCENT_COLOR = "#ff7043"  # Orange
    BG_COLOR = "#f8f9fa"  # Very light gray
    TEXT_COLOR = "#2d3748"  # Dark gray
    LIGHT_TEXT = "#718096"  # Medium gray
    BORDER_COLOR = "#e2e8f0"  # Light border
    CARD_BG = "#ffffff"  # White for cards
    
    # Fonts
    FONT_FAMILY = "Segoe UI"
    TITLE_FONT = (FONT_FAMILY, 22, "bold")
    SUBTITLE_FONT = (FONT_FAMILY, 12)
    LABEL_FONT = (FONT_FAMILY, 10, "bold")
    INPUT_FONT = (FONT_FAMILY, 10)
    CODE_FONT = "Consolas"
    
    # Dimensions
    WINDOW_SIZE = "900x800"
    MIN_WINDOW_SIZE = (800, 700)
    INPUT_WIDTH = 25
    PAD_X = 15
    PAD_Y = 10
    CARD_PADDING = 20

    # STEM subjects list
    STEM_SUBJECTS = [
        "mathematics", "maths", "further mathematics", "further maths",
        "chemistry", "physics", "biology", 
        "agricultural science", "computer science", "geography"
    ]
    
    def __init__(self, root: tk.Tk) -> None:
        """Initialize the application with the main window."""
        self.root = root
        self._setup_window()
        self._configure_styles()
        self._setup_ui()
        
    def _setup_window(self) -> None:
        """Configure the main window properties."""
        self.root.title("Avalon Exam Question Generator")
        self.root.geometry(self.WINDOW_SIZE)
        self.root.minsize(*self.MIN_WINDOW_SIZE)
        self.root.configure(bg=self.BG_COLOR)
        
        # Set window icon if available
        try:
            icon_path = os.path.join(os.path.dirname(__file__), "icon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except Exception:
            pass
        
    def _configure_styles(self) -> None:
        """Configure the ttk styles for a professional look."""
        self.style = ttk.Style()
        self.style.theme_use("clam")
        
        # Base frame style
        self.style.configure(
            "TFrame",
            background=self.BG_COLOR,
            borderwidth=0
        )
        
        # Card frame style
        self.style.configure(
            "Card.TFrame",
            background=self.CARD_BG,
            borderwidth=1,
            relief="solid",
            bordercolor=self.BORDER_COLOR
        )
        
        # Label styles
        self.style.configure(
            "TLabel",
            background=self.CARD_BG,
            foreground=self.TEXT_COLOR,
            font=self.LABEL_FONT,
            padding=0
        )
        
        self.style.configure(
            "Title.TLabel",
            font=self.TITLE_FONT,
            foreground=self.PRIMARY_COLOR,
            background=self.BG_COLOR
        )
        
        self.style.configure(
            "Subtitle.TLabel",
            font=self.SUBTITLE_FONT,
            foreground=self.LIGHT_TEXT,
            background=self.BG_COLOR
        )
        
        self.style.configure(
            "SectionTitle.TLabel",
            font=("Segoe UI", 11, "bold"),
            foreground=self.PRIMARY_COLOR,
            background=self.CARD_BG
        )
        
        # Entry styles
        self.style.configure(
            "TEntry",
            font=self.INPUT_FONT,
            foreground=self.TEXT_COLOR,
            fieldbackground="#ffffff",
            bordercolor=self.BORDER_COLOR,
            lightcolor=self.BORDER_COLOR,
            darkcolor=self.BORDER_COLOR,
            padding=8,
            relief="flat"
        )
        
        self.style.map(
            "TEntry",
            bordercolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            lightcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            darkcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ]
        )
        
        # Combobox styles
        self.style.configure(
            "TCombobox",
            font=self.INPUT_FONT,
            foreground=self.TEXT_COLOR,
            fieldbackground="#ffffff",
            bordercolor=self.BORDER_COLOR,
            lightcolor=self.BORDER_COLOR,
            darkcolor=self.BORDER_COLOR,
            padding=8,
            relief="flat"
        )
        
        self.style.map(
            "TCombobox",
            bordercolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            lightcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            darkcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ]
        )
        
        # Spinbox style
        self.style.configure(
            "TSpinbox",
            font=self.INPUT_FONT,
            foreground=self.TEXT_COLOR,
            fieldbackground="#ffffff",
            bordercolor=self.BORDER_COLOR,
            lightcolor=self.BORDER_COLOR,
            darkcolor=self.BORDER_COLOR,
            padding=8,
            relief="flat",
            arrowsize=12
        )
        
        self.style.map(
            "TSpinbox",
            bordercolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            lightcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ],
            darkcolor=[
                ('focus', self.PRIMARY_COLOR),
                ('!focus', self.BORDER_COLOR)
            ]
        )
        
        # Scrollbar style
        self.style.configure(
            "Vertical.TScrollbar",
            background=self.BORDER_COLOR,
            bordercolor=self.BORDER_COLOR,
            arrowcolor=self.TEXT_COLOR,
            troughcolor=self.BG_COLOR,
            relief="flat"
        )
        
        # LabelFrame style
        self.style.configure(
            "TLabelframe",
            background=self.CARD_BG,
            bordercolor=self.BORDER_COLOR,
            relief="solid",
            borderwidth=1
        )
        
        self.style.configure(
            "TLabelframe.Label",
            background=self.CARD_BG,
            foreground=self.PRIMARY_COLOR,
            font=("Segoe UI", 10, "bold")
        )
    
    def _setup_ui(self) -> None:
        """Create and arrange all UI components."""
        self._create_header()
        self._create_main_container()
        self._create_form_section()
        self._create_objectives_section()
        self._create_output_section()
        
    def _create_header(self) -> None:
        """Create the centered header section with title and subtitle."""
        self.header_frame = ttk.Frame(self.root, padding=(30, 20, 30, 10))
        self.header_frame.pack(fill="x")
        
        # Container frame for centered content
        center_frame = ttk.Frame(self.header_frame)
        center_frame.pack(expand=True, fill="both")
        
        # Application title (centered)
        self.title_label = ttk.Label(
            center_frame,
            text="Avalon Exam Question Generator",
            style="Title.TLabel"
        )
        self.title_label.pack(anchor="center", pady=(0, 5))
        
        # Subtitle (centered)
        self.subtitle_label = ttk.Label(
            center_frame,
            text="Create customized exam questions with AI assistance",
            style="Subtitle.TLabel"
        )
        self.subtitle_label.pack(anchor="center")
        
        # Add subtle separator
        separator = ttk.Separator(self.root, orient="horizontal")
        separator.pack(fill="x", padx=30, pady=5)
    
    def _create_main_container(self) -> None:
        """Create the main content container with scrollbar."""
        # Main container frame with padding
        self.main_container = ttk.Frame(self.root, padding=20)
        self.main_container.pack(fill="both", expand=True)
        
        # Canvas for scrolling
        self.canvas = tk.Canvas(
            self.main_container,
            bg=self.BG_COLOR,
            highlightthickness=0
        )
        self.canvas.pack(side="left", fill="both", expand=True)
        
        # Scrollbar
        self.scrollbar = ttk.Scrollbar(
            self.main_container,
            orient="vertical",
            command=self.canvas.yview
        )
        self.scrollbar.pack(side="right", fill="y")
        
        # Configure canvas scrolling
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        # Create frame inside canvas
        self.content_frame = ttk.Frame(self.canvas, padding=10)
        self.canvas.create_window((0, 0), window=self.content_frame, anchor="nw")
        
        # Mouse wheel scrolling
        self.content_frame.bind(
            "<Enter>",
            lambda _: self.content_frame.bind_all(
                "<MouseWheel>",
                lambda e: self.canvas.yview_scroll(int(-1*(e.delta/120)), "units")
            )
        )
        self.content_frame.bind(
            "<Leave>",
            lambda _: self.content_frame.unbind_all("<MouseWheel>")
        )
    
    def _create_form_section(self) -> None:
        """Create the form input section with a card layout."""
        form_frame = ttk.Frame(
            self.content_frame,
            style="Card.TFrame",
            padding=self.CARD_PADDING
        )
        form_frame.pack(fill="x", pady=(0, 20))
        
        # Section title
        section_title = ttk.Label(
            form_frame,
            text="Question Details",
            style="SectionTitle.TLabel"
        )
        section_title.grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 15))
        
        # Grid configuration
        form_frame.columnconfigure(0, weight=1)
        form_frame.columnconfigure(1, weight=1)
        form_frame.columnconfigure(2, weight=1)
        form_frame.columnconfigure(3, weight=1)
        
        # Class selection
        ttk.Label(form_frame, text="Class:").grid(
            row=1, column=0, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.class_var = tk.StringVar()
        self.class_cb = ttk.Combobox(
            form_frame,
            textvariable=self.class_var,
            state="readonly",
            values=[
                "Nursery 1", "Nursery 2", "Nursery 3",
                "Primary 1", "Primary 2", "Primary 3", "Primary 4",
                "Primary 5", "Primary 6",
                "JSS 1", "JSS 2", "JSS 3",
                "SSS 1", "SSS 2", "SSS 3"
            ],
            width=self.INPUT_WIDTH
        )
        self.class_cb.grid(
            row=1, column=1, sticky="ew", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.class_cb.current(0)
        
        # Subject entry
        ttk.Label(form_frame, text="Subject:").grid(
            row=1, column=2, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.subject_entry = ttk.Entry(
            form_frame,
            width=self.INPUT_WIDTH
        )
        self.subject_entry.grid(
            row=1, column=3, sticky="ew", padx=self.PAD_X, pady=self.PAD_Y
        )
        
        # Topic entry
        ttk.Label(form_frame, text="Topic:").grid(
            row=2, column=0, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.topic_entry = ttk.Entry(
            form_frame,
            width=self.INPUT_WIDTH
        )
        self.topic_entry.grid(
            row=2, column=1, sticky="ew", padx=self.PAD_X, pady=self.PAD_Y
        )
        
        # Question type
        ttk.Label(form_frame, text="Question Type:").grid(
            row=2, column=2, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.question_type_var = tk.StringVar()
        self.question_type_cb = ttk.Combobox(
            form_frame,
            textvariable=self.question_type_var,
            state="readonly",
            values=["Multiple Choice", "Theory", "Essay"],
            width=self.INPUT_WIDTH
        )
        self.question_type_cb.grid(
            row=2, column=3, sticky="ew", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.question_type_cb.current(0)
        
        # Number of questions
        ttk.Label(form_frame, text="Number of Questions:").grid(
            row=3, column=0, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
        self.num_questions_var = tk.IntVar(value=5)
        self.num_questions_spin = ttk.Spinbox(
            form_frame,
            from_=1,
            to=50,
            textvariable=self.num_questions_var,
            width=self.INPUT_WIDTH - 10
        )
        self.num_questions_spin.grid(
            row=3, column=1, sticky="w", padx=self.PAD_X, pady=self.PAD_Y
        )
    
    def _create_objectives_section(self) -> None:
        """Create the behavioral objectives input section."""
        objectives_frame = ttk.Frame(
            self.content_frame,
            style="Card.TFrame",
            padding=self.CARD_PADDING
        )
        objectives_frame.pack(fill="x", pady=(0, 20))
        
        # Section title
        section_title = ttk.Label(
            objectives_frame,
            text="Behavioral Objectives (minimum 3 required)",
            style="SectionTitle.TLabel"
        )
        section_title.grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 15))
        
        # Create 6 objectives in a 2x3 grid
        self.objective_entries = []
        for i in range(6):
            row = i // 3 + 1  # +1 to account for title row
            col = i % 3
            
            frame = ttk.Frame(objectives_frame, padding=(5, 5))
            frame.grid(row=row, column=col, sticky="nsew", padx=5, pady=5)
            
            ttk.Label(frame, text=f"{i + 1}.").pack(side="left", padx=(0, 5))
            entry = ttk.Entry(frame, width=30)
            entry.pack(side="left", fill="x", expand=True)
            self.objective_entries.append(entry)
            
            # Configure grid weights for resizing
            objectives_frame.columnconfigure(col, weight=1)
        
        # Generate button (centered with padding)
        btn_frame = ttk.Frame(self.content_frame)
        btn_frame.pack(fill="x", pady=(0, 20))
        
        self.generate_btn = ModernButton(
            btn_frame,
            text="Generate Questions",
            command=self.generate_questions
        )
        self.generate_btn.pack(pady=5, ipadx=20)
    
    def _create_output_section(self) -> None:
        """Create the output display section."""
        output_frame = ttk.Frame(
            self.content_frame,
            style="Card.TFrame",
            padding=self.CARD_PADDING
        )
        output_frame.pack(fill="both", expand=True)
        
        # Section title
        section_title = ttk.Label(
            output_frame,
            text="Generated Questions",
            style="SectionTitle.TLabel"
        )
        section_title.pack(anchor="w", pady=(0, 15))
        
        # Text area with scrollbars
        self.output_text = scrolledtext.ScrolledText(
            output_frame,
            height=15,
            font=tkfont.Font(family=self.CODE_FONT, size=10),
            wrap=tk.WORD,
            padx=10,
            pady=10,
            bg="#ffffff",
            fg=self.TEXT_COLOR,
            insertbackground=self.PRIMARY_COLOR,
            selectbackground=self.PRIMARY_COLOR,
            selectforeground="#ffffff",
            relief="flat",
            highlightcolor=self.PRIMARY_COLOR,
            highlightbackground=self.BORDER_COLOR,
            highlightthickness=1
        )
        self.output_text.pack(fill="both", expand=True)
        
        # Export button (centered with padding)
        btn_frame = ttk.Frame(self.content_frame)
        btn_frame.pack(fill="x", pady=(20, 0))
        
        self.export_btn = ModernButton(
            btn_frame,
            text="Export to MS Word",
            command=self.export_to_word
        )
        self.export_btn.pack(pady=5, ipadx=20)

    def is_stem_subject(self, subject: str) -> bool:
        """Checks if the given subject is a STEM subject."""
        return subject.lower() in self.STEM_SUBJECTS

    def generate_questions(self) -> None:
        """Generate exam questions based on user input."""
        try:
            if not self._validate_inputs():
                return

            prompt = self._build_prompt()
            self._display_generating_message()
            
            questions_text = self._try_generate_with_fallback(prompt)
            clean_text = self._process_ai_response(questions_text)
            
            self.output_text.config(state="normal")
            self.output_text.delete(1.0, tk.END)
            self.output_text.insert(tk.END, clean_text)
            self.output_text.config(state="disabled")

        except Exception as e:
            messagebox.showerror("Error", f"Error generating questions: {str(e)}")
    
    def _validate_inputs(self) -> bool:
        """Validate user inputs before generating questions."""
        cls = self.class_var.get().strip()
        subject = self.subject_entry.get().strip()
        topic = self.topic_entry.get().strip()
        question_type = self.question_type_var.get().strip()
        
        behavioral_objectives = [
            e.get().strip() for e in self.objective_entries if e.get().strip()
        ]
        
        if len(behavioral_objectives) < 3:
            messagebox.showerror("Error", "Please enter at least 3 behavioral objectives.")
            return False
            
        if not all([cls, subject, topic, question_type]):
            messagebox.showerror("Error", "Please fill in all required fields.")
            return False
            
        return True
    
    def _build_prompt(self) -> str:
        """Build the prompt for the AI based on user inputs."""
        cls = self.class_var.get().strip()
        subject = self.subject_entry.get().strip()
        topic = self.topic_entry.get().strip()
        question_type = self.question_type_var.get().strip()
        num_questions = self.num_questions_var.get()
        
        behavioral_objectives = [
            e.get().strip() for e in self.objective_entries if e.get().strip()
        ]
        objectives_str = "; ".join(behavioral_objectives)
        
        question_type_lower = question_type.lower()
        plural = "s" if num_questions > 1 else ""
        
        base_instructions = (
            f"Generate {num_questions} {question_type_lower}{plural} questions for {cls} {subject} "
            f"on the topic '{topic}' based on these behavioral objectives: {objectives_str}.\n"
            "Format requirements:\n"
            "- Number each question clearly\n"
            "- Use precise, unambiguous language\n"
            "- Avoid introductory or concluding remarks\n"
            "- Ensure no duplicate or redundant questions\n"
            "- Do NOT use markdown formatting like asterisks for bolding.\n"
        )
        
        if question_type == "Multiple Choice":
            base_instructions += (
                "- Include 4 choices labeled (a), (b), (c), (d)\n"
                "- All options for a single question MUST be on the same line, separated by spaces.\n" # Specific instruction for MC options
                "- Clearly indicate the correct answer for each question\n"
            )
        elif question_type == "Theory":
            base_instructions += (
                "- Questions should be open-ended requiring direct answers\n"
                "- Do not include answer choices\n"
            )
        elif question_type == "Essay":
            base_instructions += (
                "- Questions should prompt detailed explanations or discussions\n"
                "- Each question should require at least 3-5 paragraphs to answer\n"
            )
        
        # Add STEM-specific instructions if applicable
        if self.is_stem_subject(subject):
            base_instructions += (
                "\nFor STEM subjects:\n"
                "- Use proper Unicode mathematical/chemical notation (e.g., π, √, ∫, ∑, θ, ≠, ≤, ≥, H₂O, CO₂, x², y₃).\n"
                "- For equations, use proper formatting (e.g., x² + y² = z²).\n"
                "- Represent fractions clearly, e.g., 1/2 as ½ or using a fraction slash (e.g., 1⁄2).\n"
                "- Ensure formulas are correctly written with appropriate symbols and subscripts/superscripts.\n"
            )
            
        return base_instructions
    
    def _display_generating_message(self) -> None:
        """Display a message while questions are being generated."""
        self.output_text.config(state="normal")
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, "Generating questions... Please wait.\n")
        self.output_text.config(state="disabled")
        self.root.update()
    
    def _process_ai_response(self, text: str) -> str:
        """Process and clean the AI response."""
        if not text.strip():
            return "[No questions generated]"
            
        clean_text = self._clean_ai_response(text)
        clean_text = self._fix_numbering(clean_text)
        clean_text = self._remove_duplicate_questions(clean_text)

        # Process multiple choice options to be on one line
        if self.question_type_var.get() == "Multiple Choice":
            clean_text = self._consolidate_mc_options(clean_text)

        # Apply STEM formatting if applicable
        if self.is_stem_subject(self.subject_entry.get().strip()):
            clean_text = self._format_stem_content(clean_text, self.subject_entry.get().strip())
        
        return clean_text if clean_text.strip() else "[No valid questions generated]"
    
    def _call_groq_api(self, prompt: str) -> Optional[str]:
        """Call the Groq API to generate questions."""
        try:
            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
            response = client.chat.completions.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Groq API error: {e}")
            return None
    
    def _call_together_api(self, prompt: str) -> Optional[str]:
        """Call the Together.ai API to generate questions."""
        try:
            api_key = os.getenv("TOGETHER_AI_API_KEY")
            if not api_key:
                print("Together.ai API key not found.")
                return None

            url = "https://api.together.ai/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }
            json_data = {
                "model": "gpt-4o-mini", # Using a smaller model for Together.ai as a fallback
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 4000
            }
            response = requests.post(url, headers=headers, json=json_data, timeout=30)
            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"Together.ai API error: {e}")
            return None
    
    def _try_generate_with_fallback(self, prompt: str) -> str:
        """Try generating questions with primary API, fallback to secondary."""
        content = self._call_groq_api(prompt)
        if content:
            return content
            
        print("Falling back to Together.ai API")
        content = self._call_together_api(prompt)
        if content:
            return content
            
        return "[Failed to generate questions with both APIs]"
    
    def _clean_ai_response(self, text: str) -> str:
        """Remove unwanted phrases and formatting from the AI response."""
        lines = text.splitlines()
        filter_phrases = [
            "Here are", "based on", "Let me know", "Step", 
            "Here is", "meets your requirements", "Finally", 
            "In summary", "The questions are", "Answers:", "Answer:"
        ]
        filtered_lines = []
        for line in lines:
            # Remove markdown bolding asterisks
            cleaned_line = line.replace('**', '').strip()
            if not any(phrase.lower() in cleaned_line.lower() for phrase in filter_phrases):
                filtered_lines.append(cleaned_line)
        return "\n".join(filtered_lines).strip()
    
    def _fix_numbering(self, text: str) -> str:
        """Ensure consistent question numbering in the generated text."""
        lines = text.splitlines()
        result_lines = []
        q_num = 1
        question_number_pattern = r"^\s*\d+\."
        
        for line in lines:
            if line.strip() == "":
                result_lines.append(line)
                continue
                
            if line.strip().startswith(str(q_num) + "."):
                result_lines.append(line)
                q_num += 1
            elif line.strip().startswith(tuple(str(n) + "." for n in range(1, 51))):
                new_line = re.sub(question_number_pattern, f"{q_num}.", line, count=1)
                result_lines.append(new_line)
                q_num += 1
            else:
                result_lines.append(line)
                
        return "\n".join(result_lines)
    
    def _remove_duplicate_questions(self, text: str) -> str:
        """Remove duplicate questions from the generated text."""
        lines = text.splitlines()
        unique_questions = set()
        filtered_lines = []
        question_buffer = []
        # Check for MC options using (a) (b) (c) (d) pattern
        is_mc = re.search(r'\([a-d]\)', text.lower()) is not None

        def flush_question():
            if not question_buffer:
                return
            question_text = " ".join(question_buffer).strip()
            question_only = question_text
            if is_mc:
                # Remove options for uniqueness check
                question_only = re.sub(r'\s*\([a-d]\)[^\n]*', '', question_text)
            if question_only.lower() not in unique_questions:
                unique_questions.add(question_only.lower())
                filtered_lines.extend(question_buffer)
            question_buffer.clear()

        for line in lines:
            if re.match(r"^\s*\d+\.", line):
                flush_question()
                question_buffer = [line]
            elif line.strip() == "":
                flush_question()
                filtered_lines.append("")
            else:
                question_buffer.append(line)
        flush_question()

        return "\n".join(filtered_lines)

    def _consolidate_mc_options(self, text: str) -> str:
        """Consolidate multiple-choice options onto a single line per question."""
        lines = text.splitlines()
        processed_lines = []
        current_question_lines = []
        
        # Regex to match a question number at the start of a line
        question_start_pattern = re.compile(r'^\s*\d+\.\s*')
        # Regex to match an option (a), (b), (c), (d) at the start of a line
        option_start_pattern = re.compile(r'^\s*\([a-d]\)\s*')

        for line in lines:
            if question_start_pattern.match(line):
                # If we encounter a new question, process the previous one
                if current_question_lines:
                    processed_lines.extend(self._process_single_mc_question(current_question_lines))
                current_question_lines = [line]
            elif option_start_pattern.match(line) or (current_question_lines and line.strip() == ""):
                # If it's an option or an empty line following a question/option, add to current question buffer
                current_question_lines.append(line)
            else:
                # If it's a continuation of the question text or an answer, add it
                current_question_lines.append(line)
        
        # Process the last question
        if current_question_lines:
            processed_lines.extend(self._process_single_mc_question(current_question_lines))
            
        return "\n".join(processed_lines).strip()

    def _process_single_mc_question(self, question_lines: List[str]) -> List[str]:
        """Helper to process options for a single MC question."""
        if not question_lines:
            return []

        # Find the main question line (first line starting with a number)
        main_question_index = -1
        for i, line in enumerate(question_lines):
            if re.match(r'^\s*\d+\.\s*', line):
                main_question_index = i
                break
        
        if main_question_index == -1: # Should not happen if called correctly
            return question_lines

        question_text = question_lines[main_question_index].strip()
        options_buffer = []
        other_lines = []
        
        # Regex to match an option (a), (b), (c), (d)
        option_pattern = re.compile(r'^\s*\([a-d]\)\s*.*')
        
        # Iterate from the line after the main question
        for line in question_lines[main_question_index + 1:]:
            if option_pattern.match(line):
                options_buffer.append(line.strip())
            elif line.strip() == "" and options_buffer: # Allow empty lines between options if they exist
                continue
            else:
                # If we encounter a non-option line after options have started,
                # it means options have ended, and this is likely an answer or continuation.
                # Flush options and add this line to other_lines.
                if options_buffer:
                    question_text += " " + " ".join(options_buffer)
                    options_buffer = []
                other_lines.append(line)
        
        # Flush any remaining options
        if options_buffer:
            question_text += " " + " ".join(options_buffer)

        # Reconstruct the question with consolidated options
        result = [question_text]
        result.extend(other_lines) # Add any lines that followed the options (e.g., "Answer: X")
        
        return result


    def _format_stem_content(self, text: str, subject: str) -> str:
        """Post-process STEM content to ensure proper formatting for display and export."""
        subject_lower = subject.lower()
        
        # Common replacements for all STEM subjects
        replacements = {
            'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ',
            'theta': 'θ', 'pi': 'π', 'sigma': 'σ', 'omega': 'ω',
            '->': '→', '=>': '⇒', 'sqrt': '√', 'integral': '∫',
            'sum': '∑', 'product': '∏', 'infinity': '∞',
            '!=': '≠', '<=': '≤', '>=': '≥', '+-': '±',
            'deg': '°', 'lambda': 'λ', 'ohm': 'Ω',
            'approx': '≈', 'plusminus': '±', 'times': '×', 'divide': '÷'
        }
        
        for plain, unicode_char in replacements.items():
            text = text.replace(plain, unicode_char)
        
        # Subject-specific replacements
        if subject_lower in ["chemistry", "physics", "biology"]:
            text = text.replace('_2', '₂').replace('_3', '₃').replace('_4', '₄').replace('_5', '₅')
            text = text.replace('^+', '⁺').replace('^-', '⁻').replace('^2+', '²⁺').replace('^2-', '²⁻')
            
        if subject_lower == "chemistry":
            text = text.replace('<->', '⇌').replace('<=>', '⇌')
            
        # Format mathematical equations
        if subject_lower in ["mathematics", "physics", "chemistry"]:
            # Handle exponents (x^2 → x²)
            text = re.sub(r'(\w)\^(\d+)', lambda m: m.group(1) + ''.join('⁰¹²³⁴⁵⁶⁷⁸⁹'[int(d)] for d in m.group(2)), text)
            # Handle subscripts (x_1 → x₁)
            text = re.sub(r'(\w)_(\d+)', lambda m: m.group(1) + ''.join('₀₁₂₃₄₅₆₇₈₉'[int(d)] for d in m.group(2)), text)
            # Handle simple fractions like a/b to a⁄b (using fraction slash)
            text = re.sub(r'(\d+)/(\d+)', r'\1⁄\2', text)
            # Handle mixed fractions (e.g., 1 1/2 -> 1 ½)
            text = re.sub(r'(\d+)\s+(\d+)/(\d+)', r'\1 \2⁄\3', text)
            
        return text
    
    def export_to_word(self) -> None:
        """Export the generated questions to a Word document."""
        questions_raw = self.output_text.get("1.0", tk.END).strip()
        if not questions_raw:
            messagebox.showerror("Error", "No questions to export.")
            return

        cls = self.class_var.get().strip()
        subject = self.subject_entry.get().strip()
        topic = self.topic_entry.get().strip()

        if not all([cls, subject, topic]):
            messagebox.showerror(
                "Error", 
                "Please fill Class, Subject and Topic fields to export questions."
            )
            return

        # Get save location from user
        initial_filename = f"{cls}_{subject}_{topic}_questions.docx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=initial_filename,
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if not filepath:  # User cancelled
            return

        # Create and save document
        doc = Document()
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial Unicode MS' # A font that supports many Unicode characters
        font.size = Pt(11)

        doc.add_heading(f'Exam Questions for {topic}', level=1)
        doc.add_paragraph(f"Class: {cls}\nSubject: {subject}\nTopic: {topic}\n")

        questions = questions_raw.splitlines()
        for question in questions:
            if question.strip():  # Skip empty lines
                p = doc.add_paragraph()
                p.add_run(question)

        try:
            doc.save(filepath)
            messagebox.showinfo(
                "Success", 
                f"Questions successfully exported to:\n{filepath}"
            )
        except Exception as e:
            messagebox.showerror(
                "Export Error", 
                f"Failed to save document:\n{str(e)}"
            )


def main() -> None:
    """Entry point for the application."""
    root = tk.Tk()
    
    # Set Windows 10/11 theme if available
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    app = ExamQuestionGenerator(root)
    root.mainloop()


if __name__ == "__main__":
    main()
